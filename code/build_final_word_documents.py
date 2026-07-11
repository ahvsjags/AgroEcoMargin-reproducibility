"""Build polished Word deliverables from the audited v7.1 manuscript package."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "word_v7_5_prose_ready"
MAIN_MD = ROOT / "AgroEcoMargin_manuscript_v7_5_prose_revision_20260711.md"
SUPPORT_MD = ROOT / "AgroEcoMargin_v7_5_Supporting_Information_prose_revision_20260711.md"
MAIN_FIGURES = ROOT / "analysis_outputs" / "figures_v7_integrity_rebuild"
ED_FIGURES = ROOT / "analysis_outputs" / "extended_data_v7_integrity_rebuild"
RS_APPENDIX = ROOT / "analysis_outputs" / "remote_sensing_descriptive_appendix"

CITATIONS = {
    "nitrogen_response_2022": "van Grinsven et al., 2022",
    "climate_impacts_models_2021": "Jagermeyr et al., 2021",
    "climate_agriculture_adaptation_2025": "Hultgren et al., 2025",
    "kbs_lter_data_catalog_robertson_2020": "Robertson, 2020",
    "kbs_lter_public_downloads": "KBS LTER, 2026",
    "satellite_ml_2021": "Rolf et al., 2021",
    "spectral_indices_2023": "Montero et al., 2023",
    "managing_nitrogen_china_2019": "Yu et al., 2019",
    "crop_switching_climate_losses_2020": "Rising and Devineni, 2020",
    "landsat_program_2019": "Wulder et al., 2019",
    "remote_sensing_fusion_2019": "Ghamisi et al., 2019",
    "managing_nitrogen_sustainable_development_2015": "Zhang et al., 2015",
    "global_nue_trends_2014": "Lassaletta et al., 2014",
    "temperature_crop_yields_2009": "Schlenker and Roberts, 2009",
    "climate_trends_crop_production_2011": "Lobell et al., 2011",
    "climate_productivity_growth_2021": "Ortiz-Bobea et al., 2021",
    "temperature_extremes_plant_growth_2015": "Hatfield and Prueger, 2015",
    "nitrogen_under_water_stress_2024": "Drobnitch et al., 2024",
    "long_term_weather_yield_2017": "Teasdale and Cavigelli, 2017",
    "cluster_robust_inference_2015": "Cameron and Miller, 2015",
    "structured_cv_2017": "Roberts et al., 2017",
    "blockcv_2019": "Valavi et al., 2019",
    "satellite_crop_yield_mapper_2015": "Lobell et al., 2015",
    "feeding_ten_billion_boundaries_2020": "Gerten et al., 2020",
    "extreme_weather_crop_production_2016": "Lesk et al., 2016",
    "agricultural_risk_climate_2014": "Rosenzweig et al., 2014",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, before, after in (("Heading 1", 13, 16, 6), ("Heading 2", 11.5, 12, 4), ("Heading 3", 10.5, 10, 3)):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 41, 51)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(short_title)
    header_run.font.name = "Arial"
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(100, 110, 120)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def replace_citations(text: str) -> str:
    def replace(match: re.Match[str]) -> str:
        keys = re.findall(r"@([A-Za-z0-9_:-]+)", match.group(0))
        values = [CITATIONS.get(key, key) for key in keys]
        return "(" + "; ".join(values) + ")"

    return re.sub(r"\[@[^\]]+\]", replace, text)


def clean_latex(text: str) -> str:
    text = replace_citations(text)
    replacements = {
        r"\(": "", r"\)": "", r"\\": "", r"\alpha": "α", r"\beta": "β",
        r"\gamma": "γ", r"\delta": "δ", r"\tau": "τ", r"\eta": "η",
        r"\kappa": "κ", r"\varepsilon": "ε", r"\sqrt": "sqrt",
        r"\times": "x", "{": "", "}": "", "_": "_",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("ha-1", "ha⁻¹")
    return text


def add_inline(paragraph, text: str, size: float | None = None) -> None:
    text = clean_latex(text)
    token_re = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+`)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
            run.text = part[2:-2]
        elif part.startswith("*") and part.endswith("*"):
            run.italic = True
            run.text = part[1:-1]
        elif part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run.text = part[1:-1]
        else:
            run.text = part
        run.font.name = "Arial"
        if size is not None:
            run.font.size = Pt(size)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = 9360
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.autofit = True
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        table_row = table.add_row()
        cells = table_row.cells
        if row_index == 0:
            tr_pr = table_row._tr.get_or_add_trPr()
            table_header = OxmlElement("w:tblHeader")
            table_header.set(qn("w:val"), "true")
            tr_pr.append(table_header)
        for col_index in range(col_count):
            cell = cells[col_index]
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "DCE6F1")
            content = row[col_index] if col_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, content, size=8.4)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph()


def add_equation(doc: Document, raw: str) -> None:
    def math_text(parent, value: str) -> None:
        run = OxmlElement("m:r")
        text = OxmlElement("m:t")
        text.text = value
        run.append(text)
        parent.append(run)

    def math_sub(parent, base: str, subscript: str) -> None:
        node = OxmlElement("m:sSub")
        element = OxmlElement("m:e")
        sub = OxmlElement("m:sub")
        math_text(element, base)
        math_text(sub, subscript)
        node.append(element)
        node.append(sub)
        parent.append(node)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    # Render the primary model as an Office Math (OMML) object rather than
    # plain text so Word preserves its mathematical typography on submission.
    if "Y_{ipt}" in raw:
        equation = OxmlElement("m:oMath")
        math_sub(equation, "Y", "ipt")
        math_text(equation, " = ")
        math_sub(equation, "α", "p")
        math_text(equation, " + ")
        math_sub(equation, "γ", "t")
        math_text(equation, " + ")
        math_sub(equation, "δ", "c")
        math_text(equation, " ")
        math_sub(equation, "N", "ipt")
        math_text(equation, " + β ")
        math_sub(equation, "N", "ipt")
        math_sub(equation, "S", "t")
        math_text(equation, " + τ ")
        math_sub(equation, "N", "ipt")
        math_sub(equation, "T", "t")
        math_text(equation, " + ")
        math_sub(equation, "ε", "ipt")
        paragraph._p.append(equation)
        return
    run = paragraph.add_run(clean_latex(raw))
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)


def populate_markdown(doc: Document, markdown: str, include_title: bool = True) -> str:
    lines = markdown.splitlines()
    title = ""
    index = 0
    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        index = 1
        if include_title:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(10)
            run = paragraph.add_run(title)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(31, 41, 51)

    while index < len(lines):
        line = lines[index].rstrip("\n")
        if not line.strip():
            index += 1
            continue
        if line == "## Main":
            index += 1
            continue
        if line == "<!-- PAGEBREAK -->":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            index += 1
            continue
        if line.endswith("  "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            add_inline(paragraph, line.rstrip())
            index += 1
            continue
        if line.startswith("#"):
            marks = len(line) - len(line.lstrip("#"))
            heading = line[marks:].strip()
            level = 1 if marks <= 3 else 2
            doc.add_heading(clean_latex(heading), level=level)
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, parse_table(table_lines))
            continue
        if line == "\\[":
            equation_lines = []
            index += 1
            while index < len(lines) and lines[index].strip() != "\\]":
                equation_lines.append(lines[index].strip())
                index += 1
            add_equation(doc, " ".join(equation_lines))
            index += 1
            continue
        if line.startswith("---"):
            index += 1
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            nxt = lines[index].rstrip("\n")
            if not nxt.strip() or nxt.startswith("#") or nxt.startswith("|") or nxt == "\\[":
                break
            paragraph_lines.append(nxt)
            index += 1
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        add_inline(paragraph, " ".join(paragraph_lines))
    return title


def append_figures(doc: Document, directory: Path, pattern: str, heading: str, caption_prefix: str, page_break: bool = True) -> None:
    if page_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(heading, level=1)
    for index, image in enumerate(sorted(directory.glob(pattern))):
        # Keep each caption with its figure. Without an explicit break, Word can
        # leave a caption at the bottom of one page and move the image to the next.
        if index:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        number_match = re.search(r"Fig(\d+)", image.name)
        number = number_match.group(1) if number_match else image.stem
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption.paragraph_format.space_before = Pt(8)
        caption.paragraph_format.space_after = Pt(4)
        run = caption.add_run(f"{caption_prefix} {number}")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(str(image), width=Inches(6.75))


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def append_descriptive_remote_sensing(doc: Document, supporting: bool = False) -> None:
    """Add real imagery and QA context without restoring retired prediction claims."""
    image = RS_APPENDIX / (
        "Supplementary_RemoteSensing_GoldCube_QA.png" if supporting else "RemoteSensing_DataLandscape_Descriptive.png"
    )
    if not image.exists():
        return
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_landscape(section)
    heading = "Supplementary Remote-sensing Data Quality" if supporting else "Remote-sensing Data Landscape"
    doc.add_heading(heading, level=1)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    text = (
        "This descriptive figure records the image assets and quality layers used for reproducible covariate construction. "
        "It does not assert plot-level predictive skill; the strict temporal validation remains the basis for the manuscript's remote-sensing conclusion."
    )
    add_inline(intro, text)
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(image), width=Inches(9.55))
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_portrait(section)


def build_main(path: Path) -> None:
    doc = Document()
    configure_document(doc, "AgroEcoMargin")
    populate_markdown(doc, MAIN_MD.read_text(encoding="utf-8"))
    append_descriptive_remote_sensing(doc)
    append_figures(doc, MAIN_FIGURES, "Fig*.png", "Main Figures", "Fig.", page_break=False)
    doc.core_properties.title = "AgroEcoMargin Main Manuscript"
    doc.save(path)


def build_supporting(path: Path) -> None:
    doc = Document()
    configure_document(doc, "AgroEcoMargin Supporting Information")
    title = "Supporting Information"
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    paragraph.paragraph_format.space_after = Pt(10)
    populate_markdown(doc, SUPPORT_MD.read_text(encoding="utf-8"), include_title=False)
    append_descriptive_remote_sensing(doc, supporting=True)
    append_figures(doc, ED_FIGURES, "ExtendedData_Fig*.png", "Extended Data Figures", "Extended Data Fig.", page_break=False)
    doc.core_properties.title = "AgroEcoMargin Supporting Information"
    doc.save(path)


def build_combined(path: Path) -> None:
    doc = Document()
    configure_document(doc, "AgroEcoMargin Full Submission Package")
    populate_markdown(doc, MAIN_MD.read_text(encoding="utf-8"))
    append_descriptive_remote_sensing(doc)
    append_figures(doc, MAIN_FIGURES, "Fig*.png", "Main Figures", "Fig.", page_break=False)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Supporting Information")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    populate_markdown(doc, SUPPORT_MD.read_text(encoding="utf-8"), include_title=False)
    append_descriptive_remote_sensing(doc, supporting=True)
    append_figures(doc, ED_FIGURES, "ExtendedData_Fig*.png", "Extended Data Figures", "Extended Data Fig.", page_break=False)
    doc.core_properties.title = "AgroEcoMargin Full Submission Package"
    doc.save(path)


def main() -> None:
    OUT.mkdir(exist_ok=True)
    build_main(OUT / "AgroEcoMargin_v7_5_Main_Manuscript_20260711.docx")
    build_supporting(OUT / "AgroEcoMargin_v7_5_Supporting_Information_20260711.docx")
    print(f"Wrote Word deliverables to {OUT}")


if __name__ == "__main__":
    main()
