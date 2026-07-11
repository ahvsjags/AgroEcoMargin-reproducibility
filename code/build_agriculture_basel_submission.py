"""Build an Agriculture (MDPI) Article-formatted submission package.

The journal accepts free-format initial submissions, but this builder applies
the current required Article order and back-matter statements for a cleaner
submission-ready draft. Author, funding, repository and conflict metadata are
deliberately left as marked fields for author confirmation.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_final_word_documents as base


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "agriculture_basel_submission_v1"
MAIN_MD = ROOT / "AgroEcoMargin_manuscript_v7_6_submission_revision_20260711.md"
SUPPORT_MD = ROOT / "AgroEcoMargin_v7_6_Supporting_Information_submission_revision_20260711.md"
MAIN_FIGURES = ROOT / "analysis_outputs" / "figures_v7_integrity_rebuild"
ED_FIGURES = ROOT / "analysis_outputs" / "extended_data_v7_integrity_rebuild"
REMOTE_VISUALS = ROOT / "analysis_outputs" / "remote_sensing_descriptive_appendix"

# Existing bibliography order in the audited v7.5 source.
KEY_TO_SOURCE_NUMBER = {
    "nitrogen_response_2022": 1,
    "climate_impacts_models_2021": 2,
    "climate_agriculture_adaptation_2025": 3,
    "kbs_lter_data_catalog_robertson_2020": 4,
    "kbs_lter_public_downloads": 5,
    "satellite_ml_2021": 6,
    "spectral_indices_2023": 7,
    "managing_nitrogen_china_2019": 8,
    "crop_switching_climate_losses_2020": 9,
    "landsat_program_2019": 10,
    "remote_sensing_fusion_2019": 11,
    "managing_nitrogen_sustainable_development_2015": 12,
    "global_nue_trends_2014": 13,
    "temperature_crop_yields_2009": 14,
    "climate_trends_crop_production_2011": 15,
    "climate_productivity_growth_2021": 16,
    "temperature_extremes_plant_growth_2015": 17,
    "nitrogen_under_water_stress_2024": 18,
    "long_term_weather_yield_2017": 19,
    "cluster_robust_inference_2015": 20,
    "structured_cv_2017": 21,
    "blockcv_2019": 22,
    "satellite_crop_yield_mapper_2015": 23,
    "feeding_ten_billion_boundaries_2020": 24,
    "extreme_weather_crop_production_2016": 25,
    "agricultural_risk_climate_2014": 26,
    "kbs_rge_experiment": 27,
    "kbs_soil_description": 28,
}


def section(markdown: str, name: str, following: list[str]) -> str:
    start = markdown.index(f"## {name}\n") + len(f"## {name}\n")
    end = len(markdown)
    for next_name in following:
        marker = f"## {next_name}\n"
        position = markdown.find(marker, start)
        if position >= 0:
            end = min(end, position)
    return markdown[start:end].strip()


def split_references(markdown: str) -> dict[int, str]:
    raw = section(markdown, "References", [])
    matches = re.finditer(r"(?ms)^(\d+)\.\s+(.*?)(?=^\d+\.\s+|\Z)", raw)
    return {int(match.group(1)): match.group(2).strip().replace("\n", " ") for match in matches}


def numeric_citations(text: str, order: list[str]) -> str:
    def replace(match: re.Match[str]) -> str:
        values: list[int] = []
        for key in re.findall(r"@([A-Za-z0-9_:-]+)", match.group(0)):
            if key not in order:
                order.append(key)
            values.append(order.index(key) + 1)
        return "[" + ",".join(str(value) for value in sorted(set(values))) + "]"

    return re.sub(r"\[@[^\]]+\]", replace, text)


def set_hanging_indent(paragraph) -> None:
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)


def configure_mdpi(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)
    for name, size, before, after in (("Heading 1", 12, 13, 5), ("Heading 2", 11, 10, 4), ("Heading 3", 10.5, 8, 3)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_plain(doc: Document, text: str, italic: bool = False, bold: bool = False, color: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    base.add_inline(paragraph, text)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.italic = italic or run.italic
        run.bold = bold or run.bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def add_author_front_matter(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    authors = [
        ("Pengyuan Xu", "1,3,†"),
        ("Ge Zhang", "4,†"),
        ("Guang Yang", "2,3"),
        ("Moyan Li", "3,*"),
    ]
    for index, (name, markers) in enumerate(authors):
        run = paragraph.add_run(name)
        run.font.name = "Times New Roman"
        marker = paragraph.add_run(markers)
        marker.font.name = "Times New Roman"
        marker.font.superscript = True
        if index < len(authors) - 1:
            separator = paragraph.add_run("; ")
            separator.font.name = "Times New Roman"


def add_affiliation(doc: Document, marker: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    mark = paragraph.add_run(marker)
    mark.font.name = "Times New Roman"
    mark.font.superscript = True
    run = paragraph.add_run(" " + text)
    run.font.name = "Times New Roman"


def add_caption(doc: Document, number: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"Figure {number}. ")
    run.bold = True
    run.font.name = "Times New Roman"
    base.add_inline(paragraph, text)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)


def add_figure(doc: Document, number: int, caption: str) -> None:
    images = {
        1: REMOTE_VISUALS / "RemoteSensing_DataLandscape_Descriptive.png",
        2: MAIN_FIGURES / "Fig1_experiment_gradient_and_evidence_base.png",
        3: MAIN_FIGURES / "Fig2_weather_conditioned_nitrogen_response.png",
        4: MAIN_FIGURES / "Fig3_irrigation_buffering_mechanism.png",
        5: MAIN_FIGURES / "Fig4_remote_sensing_integrity_and_negative_validation.png",
    }
    image = images[number]
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_with_next = True
    picture.add_run().add_picture(str(image), width=Inches(6.55))
    add_caption(doc, number, caption)


def add_supplementary_caption(doc: Document, number: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(f"Figure S{number}. ")
    run.bold = True
    run.font.name = "Times New Roman"
    run = paragraph.add_run("Supplementary figure.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def add_supplementary_legend(doc: Document, number: int, legend: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(f"Figure S{number}. ")
    run.bold = True
    run.font.name = "Times New Roman"
    base.add_inline(paragraph, legend)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)


CAPTIONS = {
    1: "Remote-sensing data landscape. (a-c) Plot-resolved Sentinel-2 false-colour imagery, experimental plot support and plot-boundary distance at KBS. (d) Historical Landsat coverage of the KBS study area. (e-g) AOI-level Rothamsted imagery context, the KBS 24-step preharvest phenology sequence and sensor/site coverage. KBS data products are plot-resolved; Rothamsted imagery is AOI-level context only and is not used for plot-boundary inference.",
    2: "Experimental gradient and annual weather support. (a) Documented F1–F9 nitrogen schedules for corn, wheat and soybean; F1–F9 denote the nine fertilizer treatments and rates are kg N ha⁻¹. (b) Public yield-record counts and the standardized annual heat-dry index for 1999–2025. (c) Weather denominators for the rainfed and concurrent irrigated analyses. (d) Rainfed records with reported and schedule-recovered nitrogen rates.",
    3: "Association between the heat-dry index and the rainfed nitrogen response. Rainfed panel: 1,005 plot-year records, 76 plots and 27 harvest years. (a) Treatment means by heat-dry tercile. (b) Plot- and year-fixed-effect interaction estimates with year-clustered 95% confidence intervals. (c) One-year-out estimates. (d) Null distribution from 600 year-label permutations.",
    4: "Irrigation contrast in the season-specific nitrogen response. Concurrent panel: 1,674 plot-year records, 72 plots and 23 harvest years. (a) Treatment means by irrigation state and heat-dry stratum. (b) Coefficients with year-clustered 95% confidence intervals. (c) Distribution of 300 complete-year bootstrap estimates. (d) Null distribution from 600 year-label permutations.",
    5: "Plot-resolved satellite covariates and temporally separated yield prediction. (a) Preharvest Sentinel-2 normalized difference vegetation index (NDVI) within matched Main Cropping System Experiment (MCSE) plot polygons. (b) Harvest-truncated feature support by year. (c) Nested leave-one-calendar-year-out root-mean-square error (RMSE). (d) Distribution of 5,000 year-block bootstrap mean-squared-error (MSE) differences between management-plus-Landsat and management-only models.",
}


def write_content(doc: Document, text: str) -> None:
    lines = text.strip().splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line == "\\[":
            equation = []
            index += 1
            while index < len(lines) and lines[index].strip() != "\\]":
                equation.append(lines[index].strip())
                index += 1
            base.add_equation(doc, " ".join(equation))
            index += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            base.add_table(doc, base.parse_table(table_lines))
            continue
        para = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or next_line.startswith("### ") or next_line == "\\[" or next_line.startswith("|"):
                break
            para.append(next_line)
            index += 1
        add_plain(doc, " ".join(para))


def build_main() -> None:
    source = MAIN_MD.read_text(encoding="utf-8")
    title = source.splitlines()[0][2:].strip()
    abstract = section(source, "Abstract", ["Introduction"])
    intro = section(source, "Introduction", ["Results"])
    methods = section(source, "Methods", ["Figure Legends"])
    results = section(source, "Results", ["Discussion"])
    discussion = section(source, "Discussion", ["Methods"])
    references = split_references(source)
    citation_order: list[str] = []
    abstract = numeric_citations(abstract, citation_order)
    intro = numeric_citations(intro, citation_order)
    methods = numeric_citations(methods, citation_order)
    results = numeric_citations(results, citation_order)
    discussion = numeric_citations(discussion, citation_order)
    methods = methods.replace("Extended Data Fig. ", "Figure S")
    results = results.replace("Extended Data Fig. ", "Figure S")

    doc = Document()
    configure_mdpi(doc)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(9)
    run = title_p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    add_author_front_matter(doc)
    add_affiliation(doc, "1", "Department of Materials Science and Engineering, Monash University, Clayton, VIC 3800, Australia.")
    add_affiliation(doc, "2", "School of Economics and Management, China University of Mining and Technology, Xuzhou 221116, Jiangsu, China.")
    add_affiliation(doc, "3", "Hong Kong University of Science and Technology (Guangzhou), Guangzhou 510000, China.")
    add_affiliation(doc, "4", "Zhengzhou University, Zhengzhou 450001, Henan, China.")
    add_plain(doc, "† These authors contributed equally as co-first authors.")
    add_plain(doc, "* Correspondence: moyanli@hkust-gz.edu.cn (M.L.)")

    doc.add_heading("Abstract", level=1)
    add_plain(doc, abstract)
    keywords = doc.add_paragraph()
    label = keywords.add_run("Keywords: ")
    label.bold = True
    label.font.name = "Times New Roman"
    value = keywords.add_run("nitrogen response; long-term field experiment; seasonal heat and dryness; irrigation; climate adaptation; remote sensing")
    value.font.name = "Times New Roman"

    doc.add_heading("1. Introduction", level=1)
    write_content(doc, intro)
    add_figure(doc, 1, CAPTIONS[1])
    doc.add_heading("2. Materials and Methods", level=1)
    method_parts = re.split(r"(?m)^### (.+)$", methods)
    for index in range(1, len(method_parts), 2):
        heading = method_parts[index].strip()
        content = method_parts[index + 1]
        doc.add_heading(heading, level=2)
        write_content(doc, content)
        if heading == "Fertilizer-rate recovery":
            add_plain(doc, "The experimental gradient, annual weather support and treatment-rate recovery audit are summarized in Figure 2.")
            add_figure(doc, 2, CAPTIONS[2])
    doc.add_heading("3. Results", level=1)
    parts = re.split(r"(?m)^### (.+)$", results)
    for index in range(1, len(parts), 2):
        subsection = index // 2 + 1
        heading = parts[index].strip()
        content = parts[index + 1]
        doc.add_heading(f"3.{subsection}. {heading}", level=2)
        write_content(doc, content)
        add_figure(doc, subsection + 2, CAPTIONS[subsection + 2])
    doc.add_heading("4. Discussion", level=1)
    write_content(doc, discussion)
    doc.add_heading("5. Conclusions", level=1)
    add_plain(doc, "At KBS, hotter and drier seasons were associated with a lower fixed-effect estimate of marginal nitrogen response. The concurrent irrigation contrast had a positive point estimate, but its annual-resampling interval included zero. The annual diagnostics distinguish treatment replication from replication of season-level weather. In the linked plot-yield panel, preharvest Landsat and Sentinel-2 features did not provide reliable evidence of incremental out-of-year information beyond management variables.")

    doc.add_heading("Supplementary Materials", level=1)
    add_plain(doc, "The following supporting information is provided with this manuscript: Table S1: Analysis variables and provenance; Table S2: Documented F1–F9 nitrogen schedule; Table S3: Primary estimates and independent weather support; Table S4: Strict temporal satellite-yield validation; Table S5: Reproducibility map; Table S6: Crop-year support and coefficient identification; Figures S1–S10: supplementary diagnostics and validation figures; and source-data CSV files for all main and supplementary figures.")
    doc.add_heading("Author Contributions", level=1)
    add_plain(doc, "Conceptualization, P.X. and M.L.; Methodology, P.X., G.Z. and G.Y.; Software, P.X.; Validation, P.X., G.Z. and G.Y.; Formal Analysis, P.X.; Investigation, P.X., G.Z. and G.Y.; Resources, G.Y. and M.L.; Data Curation, P.X.; Writing - Original Draft Preparation, P.X.; Writing - Review and Editing, P.X., G.Z., G.Y. and M.L.; Visualization, P.X.; Supervision, M.L.; Project Administration, P.X. and M.L.; Funding Acquisition, M.L. All authors have read and agreed to the published version of the manuscript.")
    doc.add_heading("Funding", level=1)
    add_plain(doc, "This research was supported by start-up research funding awarded to M.L. by the Hong Kong University of Science and Technology (Guangzhou). No specific grant number applies.")
    doc.add_heading("Institutional Review Board Statement", level=1)
    add_plain(doc, "Not applicable.")
    doc.add_heading("Informed Consent Statement", level=1)
    add_plain(doc, "Not applicable.")
    doc.add_heading("Data Availability Statement", level=1)
    add_plain(doc, "Public KBS LTER yield and weather records are available through the KBS LTER Data Catalog. The processed analysis panels, figure-source data, bootstrap and permutation outputs, code, environment specification and documentation supporting this article are available in the fixed release at https://github.com/ahvsjags/AgroEcoMargin-reproducibility/releases/tag/v1.0.3. Provider-controlled raw data are identified by catalog accession and access condition rather than redistributed.")
    doc.add_heading("Use of Generative AI and AI-Assisted Technologies in the Writing Process", level=1)
    add_plain(doc, "During preparation of this manuscript, the authors used OpenAI Codex to assist with language editing, document formatting and reproducible-workflow code refactoring. All analytical decisions, computations, figures and manuscript claims were reviewed and verified by the authors, who take full responsibility for the content of this publication.")
    doc.add_heading("Conflicts of Interest", level=1)
    add_plain(doc, "The authors declare no conflict of interest.")

    doc.add_heading("References", level=1)
    for new_number, key in enumerate(citation_order, start=1):
        source_number = KEY_TO_SOURCE_NUMBER[key]
        paragraph = doc.add_paragraph()
        set_hanging_indent(paragraph)
        base.add_inline(paragraph, f"{new_number}. {references[source_number]}")
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
    doc.core_properties.title = "Seasonal nitrogen responses in a long-term field experiment"
    doc.save(OUT / "Agriculture_Basel_Main_Manuscript_v1.docx")
    source_order = "\n".join(f"[{index}] {key}" for index, key in enumerate(citation_order, start=1))
    (OUT / "citation_order_Agriculture_v1.txt").write_text(source_order + "\n", encoding="utf-8")


def build_supporting() -> None:
    source = SUPPORT_MD.read_text(encoding="utf-8")
    source = source.replace("# Supporting Information", "# Supplementary Materials")
    source = source.replace("Extended Data Fig. ", "Figure S")
    source = source.replace("Extended Data Figure Legends", "Supplementary Figure Legends")
    legend_start = source.find("## Supplementary Figure Legends\n")
    data_start = source.find("## Supplementary Data Package\n")
    legends: dict[int, str] = {}
    if legend_start >= 0:
        legend_end = data_start if data_start >= 0 else len(source)
        legend_block = source[legend_start + len("## Supplementary Figure Legends\n"):legend_end].strip()
        for match in re.finditer(r"(?ms)^\*\*Figure S(\d+) \| .*?\*\*\s*(.*?)(?=^\*\*Figure S\d+ \||\Z)", legend_block):
            legends[int(match.group(1))] = match.group(0).replace(f"**Figure S{match.group(1)} | ", "").replace("**", "").strip()
        source = source[:legend_start] + (source[data_start:] if data_start >= 0 else "")
    doc = Document()
    configure_mdpi(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Supplementary Materials")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    add_plain(doc, "Seasonal heat and dryness are associated with lower marginal nitrogen responses in a long-term field experiment", bold=True)
    body = source.split("\n", 1)[1].lstrip()
    body = re.sub(r"^\*\*.*?\*\*\s*\n+", "", body, count=1)
    write_content(doc, body)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Supplementary Figures", level=1)
    for image in sorted(ED_FIGURES.glob("ExtendedData_Fig*.png")):
        number = re.search(r"Fig(\d+)", image.name).group(1)
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(str(image), width=Inches(6.45))
        add_supplementary_legend(doc, int(number), legends.get(int(number), "Supplementary diagnostic."))
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(REMOTE_VISUALS / "Supplementary_RemoteSensing_GoldCube_QA.png"), width=Inches(6.45))
    add_supplementary_legend(
        doc,
        10,
        "Remote-sensing Gold-cube QA and plot support. KBS Sentinel-2 false-colour imagery and the pixel-level clear fraction, observation count, QA mode, time-gap support, layer availability and phenology-slot quality checks used to construct harvest-truncated covariates. White outlines identify the KBS plot support used for plot-resolved extraction.",
    )
    doc.core_properties.title = "Agriculture supplementary materials"
    doc.save(OUT / "Agriculture_Basel_Supplementary_Materials_v1.docx")


def write_readme() -> None:
    text = """# Agriculture (MDPI) Submission Package\n\nThis package follows the *Agriculture* Article structure: title and author front matter, abstract, keywords, Introduction, Materials and Methods, Results, Discussion, Conclusions, Supplementary Materials, Author Contributions, Funding, Data Availability Statement, Conflicts of Interest and numbered references.\n\n## Repository and upload status\n\n- Processed data, source-data CSVs, code, documentation and figure assets are publicly available in the fixed v1.0.3 release at https://github.com/ahvsjags/AgroEcoMargin-reproducibility/releases/tag/v1.0.3.\n- The main Word file includes all Figures 1-5 after their first citation; Supporting Materials is supplied as a separate Word/PDF package.\n- Confirm author-system metadata, including emails and ORCIDs where available, before final MDPI submission.\n\nThe official *Agriculture* instructions require square-bracket numeric citations, a ca. 200-word single-paragraph abstract, three to ten keywords, the Article section order, and the back-matter declarations included here.\n"""
    (OUT / "README_Agriculture_submission_v1.md").write_text(text, encoding="utf-8")


def main() -> None:
    OUT.mkdir(exist_ok=True)
    build_main()
    build_supporting()
    write_readme()
    print(f"Wrote Agriculture submission package to {OUT}")


if __name__ == "__main__":
    main()
