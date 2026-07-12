"""Generate submission-side documents for the Agriculture submission package."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "agriculture_basel_submission_v1"


def paragraph(doc: Document, text: str, *, bold: bool = False, red: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = bold
    if red:
        run.font.color.rgb = RGBColor(160, 0, 0)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85); section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = heading.add_run("Cover Letter")
    title.font.name = "Times New Roman"; title.font.size = Pt(15); title.bold = True
    paragraph(doc, "11 July 2026")
    paragraph(doc, "Editors, Agriculture")
    paragraph(doc, "Dear Editors,")
    paragraph(doc, "We submit the manuscript ‘Seasonal heat and dryness are associated with lower marginal nitrogen responses in a long-term field experiment’ for consideration as an Article in Agriculture, with Agricultural Systems and Management as the preferred section. The manuscript examines whether the marginal yield response to nitrogen changes across season-level heat and dryness in a 27-year, nine-rate field gradient, and whether a concurrent irrigation contrast moderates that pattern.")
    paragraph(doc, "The paper contributes a transparent treatment of climatic replication in long-term experiments. Across 1,005 rainfed plot-year records, hotter and drier seasons were associated with lower marginal nitrogen responses in a plot- and year-fixed-effect model. The annual-resampling interval remains wide, and the manuscript reports this uncertainty directly. The direction was retained in all 351 leave-two-year-out analyses. A concurrent rainfed-irrigated comparison provides a mechanism-consistent water-availability contrast. Separately, a nested leave-one-calendar-year-out evaluation shows that preharvest Landsat and Sentinel-2 feature blocks did not add out-of-year yield information beyond known management variables; complete outer-year diagnostics are supplied in the Supplementary Materials.")
    paragraph(doc, "The manuscript fits Agriculture because it addresses long-term agricultural-system management, nutrient response, water limitation, climate variability and reproducible assessment of digital-agriculture covariates. High-resolution RGB PNG figure files are available for independent upload if requested; original 600 dpi TIFF exports are retained for editorial request. The accompanying archive contains processed analysis tables, figure-source data, bootstrap and permutation outputs, reproducible scripts and the pinned computing environment.")
    paragraph(doc, "We confirm that neither this manuscript nor any part of its content is currently under consideration for publication elsewhere or has been published in another journal.")
    paragraph(doc, "All authors have read and approved the manuscript and agree to its submission to Agriculture.")
    paragraph(doc, "Thank you for considering this manuscript.")
    paragraph(doc, "Sincerely,")
    paragraph(doc, "Yuan Cheng, PhD\nDepartment of Materials Science and Engineering\nMonash University, Clayton, VIC 3800, Australia\nyuan.cheng@monash.edu\n\nMoyan Li, PhD\nAssistant Professor and Doctoral Supervisor\nHong Kong University of Science and Technology (Guangzhou)\nGuangzhou 510000, China\nmoyanli@hkust-gz.edu.cn")
    doc.core_properties.title = "Cover letter for Agriculture"
    doc.save(OUT / "Cover_Letter_Agriculture_v1.docx")
    print("Wrote cover letter.")


if __name__ == "__main__":
    main()
